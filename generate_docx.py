"""
Script untuk menghasilkan dokumen DOCX gabungan:
  - Isi dari final project-draft.docx (Soal 1-2)
  - Penjelasan komprehensif Soal 4c beserta penjelasan algoritma tiap file
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def set_run(run, font_name='Arial', font_size=11, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_text(doc, text, font_size=11, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, font_size=font_size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, label, description, bold_label=True):
    p = doc.add_paragraph(style='List Bullet')
    run_label = p.add_run(label)
    set_run(run_label, bold=bold_label)
    run_desc = p.add_run(description)
    set_run(run_desc)
    return p

def add_numbered(doc, label, description, bold_label=True):
    p = doc.add_paragraph(style='List Number')
    run_label = p.add_run(label)
    set_run(run_label, bold=bold_label)
    run_desc = p.add_run(description)
    set_run(run_desc)
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    # Indent
    p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = 'Arial'
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Arial'
    return table


def main():
    doc = Document()

    # ================================================================
    # HALAMAN JUDUL
    # ================================================================
    doc.add_paragraph()  # spacer
    title = doc.add_heading('FINAL PROJECT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading('ET4244 Telecommunication Optimization', 2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_text(doc, 'Paper Referensi:', font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '"Price-Based Distributed Offloading for Mobile-Edge Computing\nwith Computation Capacity Constraints"',
             font_size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ================================================================
    # DAFTAR ISI (manual)
    # ================================================================
    doc.add_heading('Daftar Isi', 1)
    toc_items = [
        'Bagian I: Background / Motivasi (Soal 1)',
        'Bagian II: Formulasi Masalah Optimasi (Soal 2)',
        'Bagian III: Uniform & Differentiated Pricing',
        'Bagian IV: Penyelesaian Soal 4c - Solve Locally (MINLP)',
        '   4c.1 Pendahuluan & Pemetaan Masalah',
        '   4c.2 Apa itu Branch and Bound (B&B)?',
        '   4c.3 Apa itu Local NLP Solver (IPOPT & MINOS)?',
        '   4c.4 Penjelasan Kode Program',
        '      - utils.py (Fungsi Utilitas)',
        '      - branch_and_bound.py (Algoritma Inti)',
        '      - compare_solvers.py (Perbandingan IPOPT vs MINOS)',
        '   4c.5 Hasil Eksekusi & Analisis',
        '      - bb_output.txt (Log Branch and Bound)',
        '      - compare_output.txt (Tabel Perbandingan)',
        '   4c.6 Kesimpulan Soal 4c',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.name = 'Arial'

    doc.add_page_break()

    # ================================================================
    # BAGIAN I: BACKGROUND (dari draft.docx)
    # ================================================================
    doc.add_heading('Bagian I: Background / Motivasi', 1)
    doc.add_heading('(Soal 1: Identify the background/motivation)', 2)

    add_text(doc, 'Perkembangan pesat aplikasi IoT (Internet of Things) mendorong kebutuhan akan komputasi real-time dengan latensi sangat rendah. Namun, perangkat IoT memiliki keterbatasan mendasar: ukuran fisik kecil, baterai terbatas, dan kemampuan komputasi yang lemah \u2014 sehingga tidak mampu memproses data intensif secara lokal dalam waktu singkat.')

    add_text(doc, 'Mobile-Edge Computing (MEC) hadir sebagai solusi dengan memindahkan (offloading) task komputasi dari perangkat pengguna ke server yang berada di tepi jaringan (network edge), bukan ke cloud terpusat yang jauh. Keunggulannya dibanding cloud konvensional:')

    add_bullet(doc, 'Latensi lebih rendah ', '(server fisik lebih dekat)')
    add_bullet(doc, 'Koneksi wireless lebih stabil ', '')
    add_bullet(doc, 'Layanan komputasi dan komunikasi lebih fleksibel', '')

    add_text(doc, 'Sebagian besar literatur sebelumnya mengasumsikan kapasitas komputasi edge cloud tidak terbatas (infinite). Padahal kenyataannya, MEC server berada di tepi jaringan dengan sumber daya terbatas \u2014 terutama saat beban kerja tinggi (intensive workload). Akibatnya:')

    add_bullet(doc, '', 'Jika terlalu banyak user offload sekaligus, kapasitas terlampaui')
    add_bullet(doc, '', 'Tidak ada mekanisme yang mengatur demand vs supply sumber daya komputasi')
    add_bullet(doc, '', 'Tidak ada insentif yang adil bagi edge cloud maupun user')

    add_text(doc, 'Paper ini mengisi gap tersebut dengan mengusulkan mekanisme pricing berbasis Stackelberg game untuk:')

    add_bullet(doc, '', 'Mengontrol jumlah data yang di-offload user agar tidak melebihi kapasitas edge cloud')
    add_bullet(doc, '', 'Memberikan insentif ekonomi yang rasional bagi edge cloud (maksimalkan revenue) dan user (minimasi latency + biaya)')
    add_bullet(doc, '', 'Dapat diimplementasikan secara terdistribusi \u2014 setiap user memutuskan secara mandiri tanpa koordinasi terpusat')

    doc.add_page_break()

    # ================================================================
    # BAGIAN II: FORMULASI MASALAH OPTIMASI (dari draft.docx)
    # ================================================================
    doc.add_heading('Bagian II: Formulasi Masalah Optimasi', 1)
    doc.add_heading('(Soal 2: Explain the optimization problem formulation)', 2)

    add_text(doc, 'Sistem terdiri dari K users dan 1 MEC server (terintegrasi dengan Base Station). Setiap user k punya R_k bits data yang harus diproses. Data bisa dibagi menjadi ell_k bits yang di-offload ke edge cloud, dan (R_k - ell_k) bits yang diproses secara lokal.')

    doc.add_heading('Komponen Waktu (Latency)', 3)
    add_text(doc, 'Waktu komputasi lokal: t_loc = (R_k - ell_k) * C_k / F_k')
    add_text(doc, 'Waktu offloading terdiri dari tiga bagian:')
    add_bullet(doc, 'Uplink: ', 't_up = ell_k / r_up_k')
    add_bullet(doc, 'Eksekusi di cloud: ', 't_cloud = ell_k * C_k / f_C')
    add_bullet(doc, 'Downlink: ', 't_down = alpha_k * ell_k / r_down_k')
    add_text(doc, 'Karena local computing dan offloading berjalan bersamaan (concurrent): t_k = max(t_loc, t_off)')

    doc.add_heading('Problem P1 \u2014 Leader Problem (Edge Cloud)', 3)
    add_text(doc, 'Edge cloud memaksimalkan revenue dari penjualan CPU cycles ke semua user:')
    add_text(doc, 'max  sum(mu_k * ell_k * C_k)  untuk semua k', italic=True)
    add_text(doc, 's.t. sum(ell_k * C_k) <= F_bar', italic=True)
    add_text(doc, 'Constraint: total CPU cycles yang dikonsumsi semua user tidak boleh melebihi kapasitas F_bar.')

    doc.add_heading('Problem P2 \u2014 Follower Problem (Setiap User k)', 3)
    add_text(doc, 'Setiap user meminimalkan latency + payment:')
    add_text(doc, 'min  t_k(ell_k) + mu_k * ell_k * C_k', italic=True)

    doc.add_heading('Solusi Optimal \u2014 Backward Induction', 3)
    add_text(doc, 'Step 1: Selesaikan P2 \u2014 Threshold Policy. User akan offload sebanyak m_k bits jika harga mu_k <= 1/F_k. Artinya offloading menguntungkan hanya jika CPU lokal lambat.')
    add_text(doc, 'Step 2: Selesaikan P1 dengan dua skema:')
    add_bullet(doc, 'Uniform Pricing: ', 'satu harga untuk semua user.')
    add_bullet(doc, 'Differentiated Pricing: ', 'harga berbeda tiap user, tereduksi menjadi Binary Knapsack.')

    doc.add_page_break()

    # ================================================================
    # BAGIAN III: UNIFORM & DIFFERENTIATED PRICING (dari draft.docx)
    # ================================================================
    doc.add_heading('Bagian III: Uniform & Differentiated Pricing', 1)

    doc.add_heading('Uniform Pricing', 2)
    add_text(doc, 'Edge cloud menetapkan satu harga mu yang sama untuk semua user. Threshold: 1/F_k merepresentasikan biaya komputasi lokal per CPU cycle user k.')
    add_bullet(doc, 'F_k besar (CPU cepat): ', '1/F_k kecil \u2014 user hanya mau offload jika harga sangat murah.')
    add_bullet(doc, 'F_k kecil (CPU lambat): ', '1/F_k besar \u2014 user toleran terhadap harga lebih tinggi karena komputasi lokal mahal.')
    add_text(doc, 'Proposition 2: Harga optimal pasti ada di himpunan {1/F_1, 1/F_2, ..., 1/F_K}. Kompleksitas pencarian: O(K log K).')

    doc.add_heading('Differentiated Pricing', 2)
    add_text(doc, 'Edge cloud menetapkan harga berbeda mu_k untuk setiap user k. Harga optimal untuk user yang dipilih offload (x_k = 1) adalah setinggi mungkin namun masih diterima user, yaitu tepat di batas: mu_k* = 1/F_k.')
    add_text(doc, 'Masalah P4 tereduksi menjadi 0-1 Knapsack Problem:')
    add_bullet(doc, 'Weight tiap item: ', 'w_k = m_k * C_k (CPU cycles yang dikonsumsi)')
    add_bullet(doc, 'Value tiap item: ', 'v_k = m_k * C_k / F_k (revenue yang dihasilkan)')
    add_text(doc, 'Diselesaikan dengan Dynamic Programming karena NP-complete.')

    doc.add_page_break()

    # ================================================================
    # BAGIAN IV: SOAL 4c - SOLVE LOCALLY (BAGIAN UTAMA YANG BARU)
    # ================================================================
    doc.add_heading('Bagian IV: Penyelesaian Soal 4c \u2014 Solve Locally (MINLP)', 1)

    # --- 4c.1 ---
    doc.add_heading('4c.1 Pendahuluan & Pemetaan Masalah', 2)
    add_text(doc, 'Soal 4c meminta kita untuk menyelesaikan masalah optimasi P4\' secara LOKAL. Artinya, kita harus menggunakan solver lokal (bukan solver global) untuk menemukan solusi optimal dari masalah yang telah diformulasikan sebagai Binary Knapsack.')

    add_text(doc, 'Masalah P4\' (Differentiated Pricing) \u2014 setelah direduksi \u2014 adalah:')
    add_code_block(doc, 'max   sum(v_k * x_k)      untuk k = 0, 1, ..., K-1')
    add_code_block(doc, 's.t.  sum(w_k * x_k) <= F_bar')
    add_code_block(doc, '      x_k dalam {0, 1}     (variabel biner / integer)')

    add_text(doc, 'Di mana:')
    add_bullet(doc, 'v_k = m_k * C_k / F_k: ', 'revenue yang didapat jika user k offload.')
    add_bullet(doc, 'w_k = m_k * C_k: ', 'siklus CPU cloud yang dikonsumsi jika user k offload.')
    add_bullet(doc, 'F_bar = 6 x 10^9 cycles/slot: ', 'kapasitas komputasi edge cloud.')
    add_bullet(doc, 'x_k: ', 'keputusan biner. Bernilai 1 jika user k diizinkan offload, 0 jika ditolak.')

    add_text(doc, 'Masalah ini tergolong MINLP (Mixed Integer Non-Linear Programming) karena keputusan x_k harus bulat (0 atau 1), yang membuat masalahnya NP-Hard. Solver optimasi biasa tidak bisa langsung menangani variabel integer secara efisien.')

    add_text(doc, 'Analogi Sederhana:', bold=True)
    add_text(doc, 'Bayangkan Anda adalah seorang satpam di acara konser musik (ini adalah Edge Cloud Server). Banyak penonton (User) ingin menitipkan tas mereka (Task Komputasi). Loker penitipan Anda berkapasitas terbatas (F_bar). Setiap tas memiliki berat (w_k = berapa ruang loker yang dihabiskan) dan biaya penitipan (v_k = pendapatan yang Anda terima). Tujuan Anda: pilih tas mana saja yang diterima agar pendapatan MAKSIMAL tanpa melebihi kapasitas loker. Setiap tas harus diterima utuh (1) atau ditolak utuh (0) \u2014 tidak bisa separuh.')

    # --- 4c.2 ---
    doc.add_heading('4c.2 Apa itu Branch and Bound (B&B)?', 2)
    add_text(doc, 'Branch and Bound adalah strategi "tebak yang pintar" untuk memecahkan masalah integer. Idenya sangat sederhana:')

    add_text(doc, 'Langkah 1 \u2014 Relaksasi (Melonggarkan Aturan)', bold=True)
    add_text(doc, 'Awalnya, kita "melonggarkan" aturan bahwa x_k harus 0 atau 1. Kita bolehkan x_k bernilai desimal (misalnya 0.3, 0.7). Ini disebut Relaksasi LP. Dengan relaksasi ini, masalah menjadi jauh lebih mudah diselesaikan oleh solver matematika (IPOPT/MINOS).')
    add_text(doc, 'Misalnya, solver memberikan jawaban: x_0 = 1.0, x_1 = 0.27, x_2 = 0.0, x_3 = 1.0, x_4 = 1.0. Di dunia nyata, kita tidak bisa menerima 27% dari tas user 1! Jadi kita perlu lanjut ke langkah berikutnya.')

    add_text(doc, 'Langkah 2 \u2014 Branching (Percabangan)', bold=True)
    add_text(doc, 'Karena x_1 = 0.27 bukan integer, kita membuat DUA SKENARIO:')
    add_bullet(doc, 'Cabang Kiri: ', 'Paksa x_1 = 0 (tolak tas user 1). Minta solver hitung ulang.')
    add_bullet(doc, 'Cabang Kanan: ', 'Paksa x_1 = 1 (terima tas user 1). Minta solver hitung ulang.')
    add_text(doc, 'Proses ini membentuk "pohon keputusan" (decision tree), di mana setiap node adalah satu skenario dengan variabel-variabel yang sudah dipaksa ke 0 atau 1.')

    add_text(doc, 'Langkah 3 \u2014 Bounding / Pruning (Pemangkasan)', bold=True)
    add_text(doc, 'Agar kita TIDAK harus mencoba semua 2^K kemungkinan (untuk K=10 berarti 1024 kombinasi!), kita "memangkas" cabang yang jelas-jelas tidak mungkin menghasilkan solusi lebih baik:')
    add_numbered(doc, 'Pangkas karena Infeasible: ', 'Jika pada suatu cabang, total bobot tas yang sudah dipaksa masuk (x_k=1) saja sudah melebihi kapasitas loker, maka cabang ini tidak mungkin feasible. Buang!')
    add_numbered(doc, 'Pangkas karena Bound: ', 'Misalkan kita sudah menemukan solusi integer dengan revenue = 10.5. Lalu di cabang baru, solver bilang "revenue MAKSIMAL (termasuk yang desimal) di cabang ini hanya 9.8". Karena 9.8 < 10.5, cabang ini mustahil mengalahkan solusi terbaik kita. Buang!')

    add_text(doc, 'Proses branching dan bounding ini terus berulang hingga seluruh pohon telah dieksplorasi atau dipangkas. Pada akhirnya, solusi integer terbaik yang ditemukan adalah SOLUSI OPTIMAL.')

    add_text(doc, 'Mengapa Metode Ini Disebut "Solve Locally"?', bold=True)
    add_text(doc, 'Disebut "lokal" karena solver yang kita gunakan di setiap node (IPOPT dan MINOS) adalah LOCAL NLP solvers. Mereka menyelesaikan sub-problem relaksasi di setiap node secara lokal (mencari titik optimal terdekat), bukan secara global. Namun karena kita membungkusnya dengan kerangka Branch and Bound yang mengeksplorasi semua kemungkinan, hasil akhirnya tetap mencapai SOLUSI OPTIMAL GLOBAL untuk masalah integer ini.')

    # --- 4c.3 ---
    doc.add_heading('4c.3 Apa itu Local NLP Solver? (IPOPT & MINOS)', 2)
    add_text(doc, 'Local NLP Solver adalah "asisten komputasi" yang bertugas memecahkan masalah matematika di mana variabelnya boleh bernilai desimal (kontinu). Mereka sangat cepat dan efisien, namun mereka TIDAK bisa memaksakan jawaban berupa 0 atau 1 murni \u2014 itulah mengapa kita butuh Branch and Bound sebagai "manajer" di atasnya.')

    add_text(doc, 'A. IPOPT (Interior Point Optimizer)', bold=True)
    add_bullet(doc, 'Metode: ', 'Interior Point Method (Metode Titik Interior)')
    add_bullet(doc, 'Analogi: ', 'Bayangkan area pencarian solusi seperti sebuah lapangan berpagar (feasible region). IPOPT mencari titik optimal dengan cara "menembus langsung lewat tengah lapangan". Ia menghitung arah terbaik dari posisinya saat ini dan bergerak lurus menuju titik optimal.')
    add_bullet(doc, 'Kekuatan: ', 'Sangat stabil dan handal untuk masalah non-linear yang sangat rumit dan besar (jutaan variabel).')
    add_bullet(doc, 'Kelemahan: ', 'Sedikit overkill untuk masalah yang sifatnya linear sederhana.')
    add_bullet(doc, 'Kompleksitas: ', 'O(n^3) per iterasi \u2014 kubik terhadap jumlah variabel.')

    add_text(doc, 'B. MINOS', bold=True)
    add_bullet(doc, 'Metode: ', 'Simplex + Reduced Gradient')
    add_bullet(doc, 'Analogi: ', 'MINOS mencari titik optimal dengan cara "menyusuri pinggiran pagar lapangan" dari sudut ke sudut. Untuk masalah yang feasible region-nya berbentuk poligon (linear), titik optimal pasti ada di salah satu sudut. Jadi MINOS sangat efisien karena langsung meloncat antar sudut.')
    add_bullet(doc, 'Kekuatan: ', 'Sangat cepat untuk masalah linear atau campuran (mixed). Memotong search space lebih agresif.')
    add_bullet(doc, 'Kelemahan: ', 'Kurang ideal untuk masalah yang sangat non-linear murni.')
    add_bullet(doc, 'Kompleksitas: ', 'O(n^2) per iterasi \u2014 kuadratik terhadap jumlah variabel.')

    add_text(doc, 'Mengapa MINOS Lebih Cepat di Soal 4c?', bold=True)
    add_text(doc, 'Sub-problem relaksasi LP pada setiap node B&B kita sebenarnya bersifat LINEAR (fungsi objektif dan constraint hanyalah penjumlahan berbobot). Untuk masalah linear, metode Simplex milik MINOS jauh lebih natural dan efisien dibandingkan Interior Point milik IPOPT. Pada K=10, MINOS hanya membutuhkan 63 node (2.35 detik) dibandingkan IPOPT yang membutuhkan 129 node (4.75 detik).')

    doc.add_page_break()

    # --- 4c.4 ---
    doc.add_heading('4c.4 Penjelasan Kode Program (File Python)', 2)

    # utils.py
    doc.add_heading('File 1: utils.py \u2014 Fungsi Utilitas & Ground Truth', 3)
    add_text(doc, 'File ini adalah "fondasi" yang menyediakan alat-alat bantu bagi program lainnya. File ini TIDAK menjalankan algoritma B&B, melainkan berisi fungsi-fungsi pendukung.')

    add_text(doc, 'Fungsi-fungsi utama di dalam utils.py:', bold=True)

    add_text(doc, '1. generate_parameters(K, seed)', bold=True)
    add_text(doc, 'Fungsi ini menghasilkan seluruh parameter simulasi untuk K user sesuai dengan Bagian IV paper:')
    add_bullet(doc, 'B = 1 MHz: ', 'Total bandwidth sistem.')
    add_bullet(doc, 'N0 = -174 dBm/Hz: ', 'Noise Power Spectral Density.')
    add_bullet(doc, 'h_k: ', 'Channel gain tiap user, uniform di [-50, -30] dB, dikonversi ke linear.')
    add_bullet(doc, 'F_k: ', 'Frekuensi CPU lokal, dipilih acak dari {0.1, 0.2, ..., 1.0} GHz.')
    add_bullet(doc, 'C_k: ', 'CPU cycles per bit, uniform di [500, 1500].')
    add_bullet(doc, 'R_k: ', 'Ukuran task data, uniform di [100, 500] KB.')
    add_bullet(doc, 'F_bar = 6 GHz: ', 'Kapasitas komputasi edge cloud.')
    add_text(doc, 'Setelah parameter dasar digenerate, fungsi ini juga menghitung kuantitas turunan:')
    add_bullet(doc, 'r_up, r_down: ', 'Kecepatan transmisi uplink dan downlink menggunakan formula Shannon.')
    add_bullet(doc, 'beta_k: ', 'Koefisien delay offloading = 1/r_up + C_k/f_C + alpha/r_down.')
    add_bullet(doc, 'm_k: ', 'Balanced offloading amount = C_k * R_k / (C_k + F_k * beta_k).')
    add_bullet(doc, 'w_k = m_k * C_k: ', 'Bobot knapsack (berapa CPU cycles yang dikonsumsi).')
    add_bullet(doc, 'v_k = m_k * C_k / F_k: ', 'Nilai knapsack (revenue yang dihasilkan).')

    add_text(doc, '2. solve_dp_knapsack(params)', bold=True)
    add_text(doc, 'Fungsi ini menyelesaikan masalah knapsack menggunakan Dynamic Programming (DP). DP adalah metode eksak yang DIJAMIN menghasilkan solusi optimal. Oleh karena itu, kita menggunakannya sebagai "kunci jawaban" (Ground Truth) untuk memverifikasi apakah algoritma Branch and Bound kita benar.')
    add_text(doc, 'Cara kerjanya: Tabel DP berukuran K x F_bar diisi dari bawah ke atas. Setiap sel dp[f] menyimpan revenue maksimum yang bisa didapat dengan kapasitas f. Setelah tabel penuh, backtrack untuk menemukan user mana saja yang terpilih.')

    add_text(doc, '3. compute_latency(params, x_opt)', bold=True)
    add_text(doc, 'Fungsi ini menghitung rata-rata latensi (waktu pemrosesan) untuk semua user berdasarkan keputusan x_k yang sudah dipilih. Latensi = max(waktu komputasi lokal, waktu offloading).')

    add_text(doc, '4. print_parameters(params)', bold=True)
    add_text(doc, 'Fungsi ini mencetak tabel parameter yang telah digenerate ke layar/file output dalam format yang rapi dan mudah dibaca.')

    doc.add_page_break()

    # branch_and_bound.py
    doc.add_heading('File 2: branch_and_bound.py \u2014 Algoritma Branch and Bound', 3)
    add_text(doc, 'File ini adalah INTI JAWABAN SOAL 4c. Di sinilah algoritma Branch and Bound diimplementasikan secara manual.')

    add_text(doc, 'Komponen-komponen utama:', bold=True)

    add_text(doc, '1. Class BranchAndBoundNode', bold=True)
    add_text(doc, 'Merepresentasikan satu node (titik keputusan) di dalam pohon B&B. Setiap node menyimpan informasi variabel mana saja yang sudah di-fix (dipaksa ke 0 atau 1) dari proses branching sebelumnya.')

    add_text(doc, '2. Fungsi build_relaxation_model(params, fixed_vars)', bold=True)
    add_text(doc, 'Membangun model relaksasi LP menggunakan library Pyomo. Model ini berisi:')
    add_bullet(doc, 'Variabel x_k dalam [0, 1]: ', 'Direlaksasi dari biner ke kontinu.')
    add_bullet(doc, 'Fungsi objektif: ', 'max sum(v_k * x_k)')
    add_bullet(doc, 'Constraint kapasitas: ', 'sum(w_k * x_k) <= F_bar')
    add_bullet(doc, 'Variabel yang di-fix: ', 'Sesuai keputusan branching dari node induk.')

    add_text(doc, '3. Fungsi solve_relaxation(params, fixed_vars, solver_name)', bold=True)
    add_text(doc, 'Menyelesaikan model relaksasi menggunakan solver yang dipilih (IPOPT atau MINOS). Jika semua variabel sudah di-fix (node daun), fungsi ini langsung mengevaluasi hasilnya tanpa memanggil solver \u2014 ini adalah optimasi penting untuk menghindari error.')

    add_text(doc, '4. Fungsi branch_and_bound(params, solver_name, verbose)', bold=True)
    add_text(doc, 'Ini adalah fungsi utama yang menjalankan seluruh algoritma B&B. Alur kerjanya:')
    add_numbered(doc, 'Inisialisasi: ', 'Buat stack (tumpukan) berisi satu node akar (root) tanpa variabel yang di-fix.')
    add_numbered(doc, 'Loop Utama: ', 'Ambil node dari stack (DFS \u2014 Depth First Search). Selesaikan relaksasi LP-nya.')
    add_numbered(doc, 'Cek Infeasible: ', 'Jika node tidak layak (kapasitas terlampaui), pangkas.')
    add_numbered(doc, 'Cek Bound: ', 'Jika nilai relaksasi <= solusi integer terbaik saat ini, pangkas.')
    add_numbered(doc, 'Cek Integralitas: ', 'Jika semua x_k sudah integer, periksa apakah ini solusi terbaik baru.')
    add_numbered(doc, 'Branching: ', 'Jika ada x_k yang pecahan (fractional), pilih variabel yang paling dekat dengan 0.5, lalu buat dua cabang anak (x_k=0 dan x_k=1).')
    add_numbered(doc, 'Ulangi: ', 'Sampai stack kosong.')
    add_text(doc, 'Setelah selesai, fungsi juga menghitung harga optimal, revenue per user, rata-rata latensi, dan persentase penggunaan kapasitas.')

    add_text(doc, '5. Fungsi print_results(result, params)', bold=True)
    add_text(doc, 'Mencetak hasil B&B dalam format tabel yang rapi, termasuk solusi optimal, metrik performa (waktu, node, pruning), dan detail per user.')

    add_text(doc, '6. Fungsi main()', bold=True)
    add_text(doc, 'Jika file ini dijalankan langsung (python branch_and_bound.py), ia akan:')
    add_bullet(doc, '', 'Mengenerate parameter untuk K=5 dan K=10 user.')
    add_bullet(doc, '', 'Menjalankan B&B dengan solver IPOPT.')
    add_bullet(doc, '', 'Memverifikasi hasilnya dengan DP (Ground Truth).')

    doc.add_page_break()

    # compare_solvers.py
    doc.add_heading('File 3: compare_solvers.py \u2014 Perbandingan IPOPT vs MINOS', 3)
    add_text(doc, 'File ini adalah script pembanding yang menjalankan algoritma B&B DKALI \u2014 sekali dengan IPOPT dan sekali dengan MINOS \u2014 lalu menyajikan perbandingan dalam bentuk tabel.')

    add_text(doc, 'Fungsi-fungsi utama:', bold=True)

    add_text(doc, '1. run_comparison(K_values, seed)', bold=True)
    add_text(doc, 'Untuk setiap nilai K (default: 5 dan 10):')
    add_bullet(doc, '', 'Generate parameter sistem.')
    add_bullet(doc, '', 'Hitung solusi DP sebagai ground truth.')
    add_bullet(doc, '', 'Jalankan B&B dengan IPOPT, catat hasilnya.')
    add_bullet(doc, '', 'Jalankan B&B dengan MINOS, catat hasilnya.')

    add_text(doc, '2. print_comparison_table(all_results)', bold=True)
    add_text(doc, 'Mencetak tabel perbandingan side-by-side yang berisi:')
    add_bullet(doc, '', 'Pendapatan optimal, rata-rata latensi, user terpilih.')
    add_bullet(doc, '', 'Waktu komputasi, jumlah node dieksplorasi.')
    add_bullet(doc, '', 'Jumlah node dipangkas (bound dan infeasible).')
    add_bullet(doc, '', 'Penggunaan kapasitas cloud.')
    add_bullet(doc, '', 'Analisis kompleksitas Big-O.')

    doc.add_page_break()

    # --- 4c.5 ---
    doc.add_heading('4c.5 Hasil Eksekusi & Analisis', 2)

    doc.add_heading('Hasil dari bb_output.txt (Log B&B dengan IPOPT)', 3)
    add_text(doc, 'File bb_output.txt berisi log detail proses Branch and Bound saat dijalankan dengan solver IPOPT. Berikut ringkasannya:')

    add_text(doc, 'Kasus K = 5 User:', bold=True)
    add_text(doc, 'Program dimulai dari Node akar (root) tanpa variabel yang di-fix. Solver IPOPT memberikan nilai relaksasi 12.558592 dan menemukan x[1] = 0.27 (fractional). Program kemudian melakukan branching pada x[1].')
    add_text(doc, 'Proses berlanjut melalui 17 node total. Beberapa node dipangkas:')
    add_bullet(doc, '3 node dipangkas karena infeasible: ', 'Kapasitas cloud terlampaui.')
    add_bullet(doc, '2 node dipangkas karena bound: ', 'Nilai relaksasi lebih kecil dari solusi terbaik.')
    add_text(doc, 'Solusi optimal akhir: User [0, 3, 4] terpilih untuk offload, dengan revenue = 10.901799 dan penggunaan kapasitas 80.7%.')

    add_text(doc, 'Kasus K = 10 User:', bold=True)
    add_text(doc, 'Masalah menjadi lebih besar dan IPOPT membutuhkan 129 node untuk menemukan solusi optimal. User [1, 2, 7] terpilih dengan revenue = 34.279394 dan penggunaan kapasitas 98.8%.')

    doc.add_heading('Hasil dari compare_output.txt (Tabel Perbandingan)', 3)
    add_text(doc, 'File compare_output.txt berisi tabel perbandingan head-to-head antara DP (Ground Truth), IPOPT, dan MINOS.')

    add_text(doc, 'Tabel Perbandingan K = 5 User:', bold=True)
    add_table(doc,
        ['Metrik', 'DP (Ground Truth)', 'IPOPT', 'MINOS'],
        [
            ['Pendapatan Optimal', '10.901799', '10.901799', '10.901799'],
            ['Rata-rata Latensi', '1.930426 s', '1.930426 s', '1.930426 s'],
            ['User Terpilih', '[0, 3, 4]', '[0, 3, 4]', '[0, 3, 4]'],
            ['Waktu Komputasi', 'N/A', '0.5737 s', '0.6241 s'],
            ['Node Dieksplorasi', 'N/A', '17', '17'],
            ['Penggunaan Kapasitas', 'N/A', '80.7%', '80.7%'],
        ]
    )

    doc.add_paragraph()  # spacer
    add_text(doc, 'Tabel Perbandingan K = 10 User:', bold=True)
    add_table(doc,
        ['Metrik', 'DP (Ground Truth)', 'IPOPT', 'MINOS'],
        [
            ['Pendapatan Optimal', '34.279394', '34.279394', '34.279394'],
            ['Rata-rata Latensi', '2.771452 s', '2.771452 s', '2.771452 s'],
            ['User Terpilih', '[1, 2, 7]', '[1, 2, 7]', '[1, 2, 7]'],
            ['Waktu Komputasi', 'N/A', '4.7544 s', '2.3522 s'],
            ['Node Dieksplorasi', 'N/A', '129', '63'],
            ['Node Dipangkas (Bound)', 'N/A', '38', '9'],
            ['Node Dipangkas (Infeasible)', 'N/A', '18', '18'],
            ['Penggunaan Kapasitas', 'N/A', '98.8%', '98.8%'],
        ]
    )

    doc.add_paragraph()  # spacer
    add_text(doc, 'Analisis Kompleksitas Big-O:', bold=True)
    add_table(doc,
        ['Metode', 'Kompleksitas Waktu', 'Kompleksitas Ruang'],
        [
            ['Branch & Bound (B&B)', 'O(2^K * T_solver)', 'O(K)'],
            ['  - IPOPT per node', 'O(n^3) per iterasi', 'O(n^2)'],
            ['  - MINOS per node', 'O(n^2) per iterasi', 'O(n^2)'],
            ['DP (Ground Truth)', 'O(K * F_bar/delta)', 'O(K * F_bar/delta)'],
        ]
    )

    doc.add_page_break()

    # --- 4c.6 ---
    doc.add_heading('4c.6 Kesimpulan Soal 4c', 2)

    add_text(doc, '1. Akurasi Solusi:', bold=True)
    add_text(doc, 'Algoritma Branch and Bound manual yang kami implementasikan, dengan relaksasi NLP menggunakan IPOPT maupun MINOS, berhasil mencapai SOLUSI OPTIMAL GLOBAL. Hal ini dibuktikan dengan kecocokan 100% terhadap solusi Dynamic Programming (Ground Truth) untuk semua skenario uji (K=5 dan K=10).')

    add_text(doc, '2. Perbandingan Kinerja Solver:', bold=True)
    add_text(doc, 'Pada skala kecil (K=5), kedua solver memiliki performa yang hampir identik. Namun pada skala lebih besar (K=10), MINOS menunjukkan keunggulan signifikan:')
    add_bullet(doc, 'Waktu komputasi: ', 'MINOS 50% lebih cepat (2.35 vs 4.75 detik).')
    add_bullet(doc, 'Efisiensi pencarian: ', 'MINOS hanya mengeksplorasi 63 node vs 129 node pada IPOPT.')
    add_bullet(doc, 'Alasan: ', 'Sub-problem relaksasi LP pada Knapsack bersifat LINEAR, sehingga metode Simplex (MINOS) jauh lebih efisien daripada Interior Point (IPOPT).')

    add_text(doc, '3. Validitas Pendekatan "Solve Locally":', bold=True)
    add_text(doc, 'Meskipun IPOPT dan MINOS adalah solver LOKAL (mereka hanya mencari solusi di sekitar titik terdekat), namun karena dibungkus dalam kerangka Branch and Bound yang mengeksplorasi seluruh search space secara sistematis, hasil akhirnya tetap merupakan OPTIMUM GLOBAL untuk masalah integer ini. Ini membuktikan bahwa pendekatan "solve locally + B&B" adalah strategi yang valid dan efektif untuk menyelesaikan MINLP.')

    add_text(doc, '4. Kompleksitas:', bold=True)
    add_text(doc, 'Secara teori, kompleksitas waktu B&B adalah O(2^K) di skenario terburuk (eksponensial). Namun dalam praktiknya, pruning (pemangkasan) secara dramatis mengurangi jumlah node yang perlu dieksplorasi. Untuk K=10 dengan 1024 kemungkinan, MINOS hanya perlu mengecek 63 node (6.2% dari total kemungkinan).')

    # Save
    output_path = 'Final_Project_Gabungan_Soal_4c.docx'
    doc.save(output_path)
    print(f"Dokumen berhasil disimpan: {output_path}")

if __name__ == "__main__":
    main()
